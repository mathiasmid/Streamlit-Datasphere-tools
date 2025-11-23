"""
Documentation builder for SAP Datasphere lineage.

Generates comprehensive documentation from lineage and CSN definitions,
including field-level mappings and transformation logic.
"""

import logging
from typing import Dict, Any, List, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from .models import LineageTree, LineageNode, CSNDefinition, CSNElement, AppConfig
from .api_client import DataspherAPIClient
from .db_client import HANAClient
from .lineage import LineageAnalyzer

logger = logging.getLogger(__name__)


class DocumentationBuilder:
    """
    Builds comprehensive documentation from lineage and object definitions.
    """

    def __init__(self, config: AppConfig):
        """
        Initialize documentation builder.

        Args:
            config: Application configuration
        """
        self.config = config
        self.api_client = DataspherAPIClient(config)
        self.db_client = HANAClient(config)

    def build_lineage_documentation(
        self,
        lineage_tree: LineageTree,
        root_object_name: str,
        include_field_mappings: bool = True,
        include_transformations: bool = True,
        transactional_only: bool = False
    ) -> Document:
        """
        Build complete documentation for lineage.

        Args:
            lineage_tree: Lineage tree to document
            root_object_name: Name of root object
            include_field_mappings: Include field-level mappings
            include_transformations: Include transformation logic
            transactional_only: Document only transactional objects

        Returns:
            Word Document object
        """
        logger.info(f"Building documentation for {root_object_name}")

        # Filter if needed
        if transactional_only:
            filtered_tree = lineage_tree.get_transactional_lineage()
            if filtered_tree:
                lineage_tree = filtered_tree

        # Create document
        doc = Document()

        # Add title and metadata
        self._add_title_page(doc, root_object_name, lineage_tree)

        # Add table of contents placeholder
        self._add_toc_placeholder(doc)

        # Add overview section
        self._add_overview_section(doc, lineage_tree, root_object_name)

        # Add detailed object documentation
        self._add_objects_section(
            doc,
            lineage_tree,
            include_field_mappings,
            include_transformations
        )

        # Add field mapping summary
        if include_field_mappings:
            self._add_field_mapping_section(doc, lineage_tree)

        # Add appendix
        self._add_appendix(doc, lineage_tree)

        logger.info("Documentation built successfully")
        return doc

    def _add_title_page(self, doc: Document, object_name: str, lineage_tree: LineageTree):
        """Add title page to document."""
        # Title
        title = doc.add_heading(f"Data Lineage Documentation", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_heading(object_name, 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph()
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Light Grid Accent 1'

        metadata_table.rows[0].cells[0].text = "Generated On"
        metadata_table.rows[0].cells[1].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        metadata_table.rows[1].cells[0].text = "Lineage Fetched"
        metadata_table.rows[1].cells[1].text = lineage_tree.fetched_at.strftime("%Y-%m-%d %H:%M:%S")

        metadata_table.rows[2].cells[0].text = "Total Objects"
        metadata_table.rows[2].cells[1].text = str(lineage_tree.count_objects())

        transactional_count = sum(1 for obj in lineage_tree.get_all_objects() if obj.is_transactional())
        metadata_table.rows[3].cells[0].text = "Transactional Objects"
        metadata_table.rows[3].cells[1].text = str(transactional_count)

        doc.add_page_break()

    def _add_toc_placeholder(self, doc: Document):
        """Add table of contents placeholder."""
        doc.add_heading("Table of Contents", 1)
        doc.add_paragraph(
            "Note: In Microsoft Word, right-click here and select 'Update Field' to generate the table of contents."
        )
        doc.add_page_break()

    def _add_overview_section(self, doc: Document, lineage_tree: LineageTree, object_name: str):
        """Add overview section with lineage summary."""
        doc.add_heading("1. Overview", 1)

        doc.add_heading("1.1 Purpose", 2)
        doc.add_paragraph(
            f"This document provides comprehensive lineage documentation for {object_name}, "
            f"including all dependencies, data flows, and field-level mappings."
        )

        doc.add_heading("1.2 Lineage Summary", 2)

        # Analyze lineage
        analyzer = LineageAnalyzer(self.config)
        analysis = analyzer.analyze_lineage(lineage_tree)

        # Summary table
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Light Grid Accent 1'

        summary_table.rows[0].cells[0].text = "Total Objects in Lineage"
        summary_table.rows[0].cells[1].text = str(analysis['total_objects'])

        summary_table.rows[1].cells[0].text = "Transactional Objects"
        summary_table.rows[1].cells[1].text = str(analysis['transactional_objects'])

        summary_table.rows[2].cells[0].text = "Non-Transactional Objects"
        summary_table.rows[2].cells[1].text = str(analysis['non_transactional_objects'])

        summary_table.rows[3].cells[0].text = "Maximum Dependency Depth"
        summary_table.rows[3].cells[1].text = str(analysis['max_dependency_depth'])

        summary_table.rows[4].cells[0].text = "Source Objects (Leaf Nodes)"
        summary_table.rows[4].cells[1].text = str(analysis['source_object_count'])

        # Object type distribution
        doc.add_heading("1.3 Object Type Distribution", 2)

        type_table = doc.add_table(rows=len(analysis['object_type_counts']) + 1, cols=2)
        type_table.style = 'Light Grid Accent 1'

        type_table.rows[0].cells[0].text = "Object Type"
        type_table.rows[0].cells[1].text = "Count"

        for idx, (obj_type, count) in enumerate(analysis['object_type_counts'].items(), 1):
            type_table.rows[idx].cells[0].text = obj_type
            type_table.rows[idx].cells[1].text = str(count)

        doc.add_page_break()

    def _add_objects_section(
        self,
        doc: Document,
        lineage_tree: LineageTree,
        include_fields: bool,
        include_transformations: bool
    ):
        """Add detailed section for each object."""
        doc.add_heading("2. Object Details", 1)

        all_objects = lineage_tree.get_all_objects()

        for idx, obj in enumerate(all_objects, 1):
            doc.add_heading(f"2.{idx} {obj.name}", 2)

            # Basic info
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Light Grid Accent 1'

            info_table.rows[0].cells[0].text = "Qualified Name"
            info_table.rows[0].cells[1].text = obj.qualified_name

            info_table.rows[1].cells[0].text = "Object Type"
            info_table.rows[1].cells[1].text = obj.kind

            info_table.rows[2].cells[0].text = "Transactional"
            info_table.rows[2].cells[1].text = "Yes" if obj.is_transactional() else "No"

            info_table.rows[3].cells[0].text = "Dependencies"
            info_table.rows[3].cells[1].text = str(len(obj.dependencies))

            # Get field information if available
            if include_fields:
                field_data = None
                space_id = None

                try:
                    # Extract space from qualified name or folder
                    space_id = self._extract_space_id(obj.qualified_name, obj.folder_id)

                    logger.info(f"Attempting to get field info for {obj.qualified_name}, space_id: {space_id}")

                    if space_id:
                        # Try CSN first (more detailed metadata)
                        try:
                            logger.info(f"Querying CSN for {obj.qualified_name} in space {space_id}")
                            csn_def = self.db_client.get_object_csn(space_id, obj.qualified_name)
                            if csn_def and csn_def.elements:
                                field_data = ('csn', csn_def.elements)
                                logger.info(f"Successfully retrieved {len(csn_def.elements)} fields from CSN")
                        except Exception as csn_error:
                            logger.warning(f"CSN query failed for {obj.qualified_name}: {csn_error}")
                            # If CSN fails (e.g., $$DEPLOY_ARTIFACTS$$ not available), try M_CS_COLUMNS
                            if "$$DEPLOY_ARTIFACTS$$" in str(csn_error) or "invalid table name" in str(csn_error).lower():
                                logger.info(f"Trying M_CS_COLUMNS as fallback for {obj.qualified_name}")
                                try:
                                    columns = self.db_client.get_table_columns(space_id, obj.qualified_name)
                                    if columns:
                                        field_data = ('columns', columns)
                                        logger.info(f"Successfully retrieved {len(columns)} columns from M_CS_COLUMNS")
                                except Exception as col_error:
                                    logger.warning(f"M_CS_COLUMNS also failed for {obj.qualified_name}: {col_error}")
                            else:
                                raise csn_error

                        if field_data:
                            data_type, data = field_data
                            doc.add_heading(f"2.{idx}.1 Fields", 3)

                            if data_type == 'csn':
                                # CSN data with full metadata
                                field_table = doc.add_table(rows=len(data) + 1, cols=6)
                                field_table.style = 'Light Grid Accent 1'

                                # Header
                                header_cells = field_table.rows[0].cells
                                header_cells[0].text = "Key"
                                header_cells[1].text = "Required"
                                header_cells[2].text = "Field Name"
                                header_cells[3].text = "Label"
                                header_cells[4].text = "Type"
                                header_cells[5].text = "Length"

                                # Data rows
                                for field_idx, element in enumerate(data, 1):
                                    cells = field_table.rows[field_idx].cells
                                    cells[0].text = "✓" if element.key else ""
                                    cells[1].text = "✓" if element.not_null else ""
                                    cells[2].text = element.technical_name
                                    cells[3].text = element.label or "-"
                                    cells[4].text = element.type
                                    cells[5].text = str(element.length) if element.length else "-"
                            else:
                                # Column data from M_CS_COLUMNS
                                field_table = doc.add_table(rows=len(data) + 1, cols=5)
                                field_table.style = 'Light Grid Accent 1'

                                # Header
                                header_cells = field_table.rows[0].cells
                                header_cells[0].text = "Nullable"
                                header_cells[1].text = "Field Name"
                                header_cells[2].text = "Type"
                                header_cells[3].text = "Length"
                                header_cells[4].text = "Scale"

                                # Data rows
                                for field_idx, col in enumerate(data, 1):
                                    cells = field_table.rows[field_idx].cells
                                    cells[0].text = "Yes" if col.get('IS_NULLABLE') == 'TRUE' else "No"
                                    cells[1].text = col.get('COLUMN_NAME', '-')
                                    cells[2].text = col.get('DATA_TYPE_NAME', '-')
                                    cells[3].text = str(col.get('LENGTH', '-'))
                                    cells[4].text = str(col.get('SCALE', '-')) if col.get('SCALE') else "-"
                        else:
                            logger.debug(f"No field information found for {obj.qualified_name}")
                    else:
                        logger.debug(f"Could not extract space_id for {obj.qualified_name}, skipping field information")

                except Exception as e:
                    logger.warning(f"Failed to fetch field information for {obj.qualified_name}: {e}")
                    # Only show error message if it's not a known table issue
                    if "$$DEPLOY_ARTIFACTS$$" not in str(e) and "invalid table name" not in str(e).lower():
                        doc.add_paragraph(f"Note: Field information not available ({str(e)[:100]})")

            doc.add_paragraph()  # Spacing

    def _add_field_mapping_section(self, doc: Document, lineage_tree: LineageTree):
        """Add field mapping summary section."""
        doc.add_page_break()
        doc.add_heading("3. Field Mappings", 1)

        doc.add_paragraph(
            "This section shows field-level mappings across the lineage. "
            "For each transformation, source fields are mapped to target fields."
        )

        # Get transactional objects only (flows)
        transactional_objects = [
            obj for obj in lineage_tree.get_all_objects()
            if obj.is_transactional()
        ]

        if not transactional_objects:
            doc.add_paragraph("No transactional objects found in lineage.")
            return

        for idx, obj in enumerate(transactional_objects, 1):
            doc.add_heading(f"3.{idx} {obj.name}", 2)

            # For transformation flows, show source → target mapping
            if 'transformationflow' in obj.kind.lower():
                doc.add_paragraph(f"Type: Transformation Flow")

                # List source dependencies
                if obj.dependencies:
                    doc.add_paragraph("Source Objects:")
                    for dep in obj.dependencies:
                        if dep.dependency_type and 'source' in dep.dependency_type.lower():
                            doc.add_paragraph(f"  • {dep.qualified_name}", style='List Bullet')

            elif 'replicationflow' in obj.kind.lower():
                doc.add_paragraph(f"Type: Replication Flow")
                doc.add_paragraph("Direct data replication from source to target.")

            doc.add_paragraph()  # Spacing

    def _add_appendix(self, doc: Document, lineage_tree: LineageTree):
        """Add appendix with additional information."""
        doc.add_page_break()
        doc.add_heading("Appendix", 1)

        doc.add_heading("A. Complete Object List", 2)

        # Flat list of all objects
        all_objects = lineage_tree.get_all_objects()

        obj_table = doc.add_table(rows=len(all_objects) + 1, cols=4)
        obj_table.style = 'Light Grid Accent 1'

        # Header
        obj_table.rows[0].cells[0].text = "#"
        obj_table.rows[0].cells[1].text = "Object Name"
        obj_table.rows[0].cells[2].text = "Type"
        obj_table.rows[0].cells[3].text = "Transactional"

        # Data
        for idx, obj in enumerate(all_objects, 1):
            obj_table.rows[idx].cells[0].text = str(idx)
            obj_table.rows[idx].cells[1].text = obj.qualified_name
            obj_table.rows[idx].cells[2].text = obj.kind
            obj_table.rows[idx].cells[3].text = "Yes" if obj.is_transactional() else "No"

        # Glossary
        doc.add_heading("B. Glossary", 2)

        glossary_terms = [
            ("Lineage", "The path of data from source systems through transformations to final target objects."),
            ("Transactional Object", "Objects that modify or move data (flows, replications, transformations)."),
            ("Non-Transactional Object", "Objects that reference data without modifying it (views, associations)."),
            ("Replication Flow", "SAP Datasphere object that replicates data from source to target."),
            ("Transformation Flow", "SAP Datasphere object that transforms data using SQL or graphical logic."),
            ("Dependency", "Relationship between objects where one object uses or references another."),
            ("CSN", "Core Schema Notation - SAP's metadata format containing field definitions.")
        ]

        for term, definition in glossary_terms:
            p = doc.add_paragraph()
            p.add_run(f"{term}: ").bold = True
            p.add_run(definition)

    def _extract_space_id(self, qualified_name: str, folder_id: str) -> Optional[str]:
        """
        Extract space ID from qualified name or folder ID using the same logic as documentation_helper.

        Args:
            qualified_name: Object qualified name
            folder_id: Folder/space ID

        Returns:
            Space ID or None
        """
        # Define prefix to space mapping (same as documentation_helper.py)
        prefix_to_space = {
            '01_ACQ': '01_ACQUISITION',
            '02_DWH': '02_DATAWAREHOUSE',
            '03_SAL': '03_SALES',
            '03_FIN': '03_FINANCE',
            '04_PBI': '04_POWERBI'
        }

        # Check each prefix
        for prefix, space in prefix_to_space.items():
            if qualified_name.startswith(prefix):
                return space

        # No matching prefix found
        return None


def save_documentation(doc: Document, filename: str) -> bool:
    """
    Save documentation to file.

    Args:
        doc: Document to save
        filename: Output filename

    Returns:
        True if successful
    """
    try:
        doc.save(filename)
        logger.info(f"Documentation saved to {filename}")
        return True

    except Exception as e:
        logger.error(f"Failed to save documentation: {e}")
        return False
